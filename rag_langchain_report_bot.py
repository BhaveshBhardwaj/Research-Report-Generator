import streamlit as st
import os
import time
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.chains.retrieval import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
import re
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import io
from fpdf import FPDF
import unicodedata
import docx

def deep_clean_text(text):
    """
    Performs a deep cleaning of the text to remove all non-printable and problematic
    Unicode characters that can cause crashes in PDF generation libraries.
    """
    if not isinstance(text, str):
        return ""
    text = unicodedata.normalize('NFKD', text)
    cleaned_chars = []
    for char in text:
        category = unicodedata.category(char)
        if category.startswith(('L', 'N', 'P', 'S', 'Z')):
            cleaned_chars.append(char)
    return "".join(cleaned_chars)

class PDF(FPDF):
    def __init__(self, font_family='DejaVu', topic=""):
        super().__init__()
        self.font_family = font_family
        self.topic = topic

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font(self.font_family, '', 9)
        self.cell(0, 10, self.topic, 0, 0, 'L')
        self.ln(10)

    def footer(self):
        if self.page_no() == 1:
            return
        self.set_y(-15)
        self.set_font(self.font_family, 'I', 8)
        self.cell(0, 10, f'Page {self.page_no() - 1}', 0, 0, 'C')

def create_download_pdf(report_content, topic="AI Research Report"):
    """
    Generates a well-formatted, Unicode-aware PDF with a title page, headers,
    footers, and proper page breaks.
    """
    cleaned_topic = deep_clean_text(topic)
    pdf = PDF(topic=cleaned_topic)
    
    try:
        pdf.add_font('DejaVu', '', 'dejavu-sans/DejaVuSans.ttf', uni=True)
        pdf.add_font('DejaVu', 'B', 'dejavu-sans/DejaVuSans-Bold.ttf', uni=True)
        pdf.add_font('DejaVu', 'I', 'dejavu-sans/DejaVuSans-Oblique.ttf', uni=True)
    except RuntimeError:
        st.error("CRITICAL: DejaVu font files not found. Please ensure 'dejavu-sans' folder exists with the required .ttf files.")
        return None

    pdf.add_page()
    pdf.set_font('DejaVu', 'B', 24)
    pdf.ln(60)
    pdf.multi_cell(0, 15, cleaned_topic, 0, 'C')
    pdf.ln(10)
    pdf.set_font('DejaVu', '', 14)
    pdf.multi_cell(0, 10, "An AI-Generated Research Analysis", 0, 'C')
    pdf.set_y(-30)
    pdf.set_font('DejaVu', '', 10)
    pdf.multi_cell(0, 10, f"Generated on: {time.strftime('%Y-%m-%d')}", 0, 'C')

    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=25)
    
    report_body = re.sub(r'#.*?\n+', '', report_content, count=1).strip()
    blocks = re.split(r'\n\s*\n', report_body)

    for block in blocks:
        block = block.strip()
        if not block: continue
        cleaned_block = deep_clean_text(block)
        is_heading = cleaned_block.startswith('## ')
        
        if is_heading and pdf.get_y() > (297 - 50):
             pdf.add_page()

        if is_heading:
            pdf.ln(8)
            pdf.set_font('DejaVu', 'B', 16)
            pdf.multi_cell(0, 8, cleaned_block.replace('## ', ''), 0, 'L')
            pdf.ln(4)
        else:
            pdf.set_font('DejaVu', '', 12)
            pdf.multi_cell(0, 7, cleaned_block, 0, 'L')
            pdf.ln(3)

    return bytes(pdf.output())

def create_download_docx(report_content, topic="AI Research Report"):
    """
    Generates a .docx file from the report string and returns its bytes.
    """
    document = docx.Document()
    cleaned_topic = deep_clean_text(topic)
    
    document.add_heading(cleaned_topic, level=0)
    
    report_body = re.sub(r'#.*?\n+', '', report_content, count=1).strip()
    blocks = re.split(r'\n\s*\n', report_body)

    for block in blocks:
        block = block.strip()
        if not block: continue
        
        cleaned_block = deep_clean_text(block)
        
        if cleaned_block.startswith('## '):
            document.add_heading(cleaned_block.replace('## ', ''), level=2)
        else:
            document.add_paragraph(cleaned_block)
            
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def process_pdf(pdf_files):
    if not pdf_files: return []
    full_text = ""
    st.write("Extracting text and performing OCR on images...")
    progress_bar = st.progress(0)
    for i, pdf_file in enumerate(pdf_files):
        try:
            doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                full_text += page.get_text()
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    try:
                        image = Image.open(io.BytesIO(image_bytes))
                        ocr_text = pytesseract.image_to_string(image)
                        if ocr_text.strip():
                            full_text += f"\n[OCR Text from Image on page {page_num + 1}]:\n{ocr_text}\n"
                    except Exception:
                        pass
            doc.close()
        except Exception as e:
            st.error(f"Error processing {pdf_file.name}: {e}")
        progress_bar.progress((i + 1) / len(pdf_files))
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200, length_function=len)
    return text_splitter.split_text(text=full_text)

def get_vector_store(text_chunks):
    if not text_chunks: return None
    try:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        return FAISS.from_texts(text_chunks, embedding=embeddings)
    except Exception as e:
        st.error(f"Failed to create vector store: {e}")
        return None

def create_rag_chain(vector_store, groq_api_key, prompt_template):
    llm = ChatGroq(groq_api_key=groq_api_key, model_name="meta-llama/llama-4-scout-17b-16e-instruct")
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "input"])
    retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 5})
    combine_docs_chain = create_stuff_documents_chain(llm, prompt)
    return create_retrieval_chain(retriever, combine_docs_chain)

def main():
    st.set_page_config(page_title="Advanced RAG Report Generator", layout="wide")
    st.title("📚 Advanced AI Research Report Generator")
    st.markdown("""
    Welcome! This tool generates detailed research reports with PDF download and OCR capabilities.
    1.  **RAG Mode:** Upload PDFs to generate a report from their content.
    2.  **Generative Mode:** Leave the file uploader empty and provide a topic to generate a report from scratch.
    """)

    if 'report_generated' not in st.session_state:
        st.session_state.report_generated = False
        st.session_state.final_report = ""
        st.session_state.report_topic = ""

    with st.sidebar:
        st.header("⚙️ Configuration")
        groq_api_key = st.text_input("Enter your Groq API Key:", type="password")
        page_count = st.number_input("Desired Number of Pages:", min_value=1, max_value=200, value=10, step=1)
        user_instructions = st.text_area("Optional Instructions:", height=150, placeholder="e.g., 'Focus on financial analysis.'")

    st.subheader("Step 1: Provide Content Source")
    report_topic = st.text_input("Enter a Topic (if not uploading files)", placeholder="e.g., 'The Future of Renewable Energy'")
    uploaded_files = st.file_uploader("OR Upload PDF Document(s) with OCR (Optional)", type="pdf", accept_multiple_files=True)

    if st.button("Generate Detailed Report"):
        if not groq_api_key: st.error("A Groq API Key is required."); return
        if not uploaded_files and not report_topic: st.error("Please upload PDFs or enter a topic."); return
        
        final_report = ""
        current_topic = ""

        if uploaded_files:
            with st.spinner("Processing documents with OCR and generating RAG report..."):
                text_chunks = process_pdf(uploaded_files)
                if text_chunks:
                    vector_store = get_vector_store(text_chunks)
                    if vector_store:
                        instruction_addition = f"USER INSTRUCTIONS: \"{user_instructions}\"" if user_instructions else ""
                        outline_prompt = f"""Based on the provided context, generate a detailed table of contents for a research report of about {page_count} pages. {instruction_addition} CONTEXT: {{context}} QUERY: {{input}}"""
                        outline_chain = create_rag_chain(vector_store, groq_api_key, outline_prompt)
                        toc = outline_chain.invoke({"input": "Generate a table of contents."})['answer']
                        with st.expander("Generated Table of Contents", expanded=True): st.markdown(toc)
                        section_titles = re.findall(r'^\s*[\d\w]+\..*$', toc, re.MULTILINE)
                        if not section_titles: section_titles = [line.strip() for line in toc.split('\n') if line.strip() and len(line.strip()) > 5]
                        full_report_list = []
                        words_per_section = (page_count * 400) // len(section_titles) if section_titles else 500
                        section_prompt = f"""As a research analyst, write a comprehensive section on: "{{input}}". Write about {words_per_section} words. Use the CONTEXT to elaborate. {instruction_addition} DO NOT write a title. CONTEXT: {{context}} QUERY: {{input}}"""
                        section_chain = create_rag_chain(vector_store, groq_api_key, section_prompt)
                        progress_bar = st.progress(0, text="Generating report sections...")
                        for i, title in enumerate(section_titles):
                            clean_title = re.sub(r'^\s*[\d\w]+\.\s*', '', title).strip()
                            st.write(f"-> Generating: {clean_title}")
                            response = section_chain.invoke({"input": clean_title})
                            full_report_list.append(f"## {title}\n\n{response['answer']}")
                            progress_bar.progress((i + 1) / len(section_titles), text=f"Generated: {clean_title}")
                        final_report = "# Your In-Depth Research Report\n\n" + "\n\n".join(full_report_list)
                        current_topic = "RAG-Generated Research Report"
        else:
            with st.spinner(f"Generating a new report on '{report_topic}'..."):
                llm = ChatGroq(groq_api_key=groq_api_key, model_name="meta-llama/llama-4-scout-17b-16e-instruct")
                instruction_addition = f"USER INSTRUCTIONS: \"{user_instructions}\"" if user_instructions else ""
                outline_prompt = PromptTemplate.from_template(f"Generate a detailed table of contents for a report of about {page_count} pages on the topic: '{{topic}}'. {instruction_addition}")
                outline_chain = outline_prompt | llm
                toc = outline_chain.invoke({"topic": report_topic}).content
                with st.expander("Generated Table of Contents", expanded=True): st.markdown(toc)
                section_titles = re.findall(r'^\s*[\d\w]+\..*$', toc, re.MULTILINE)
                if not section_titles: section_titles = [line.strip() for line in toc.split('\n') if line.strip() and len(line.strip()) > 5]
                full_report_list = []
                words_per_section = (page_count * 400) // len(section_titles) if section_titles else 500
                section_prompt = PromptTemplate.from_template(f"As a research analyst, write a comprehensive section on: '{{section}}'. The main report topic is '{report_topic}'. Write about {words_per_section} words. {instruction_addition} DO NOT write a title.")
                section_chain = section_prompt | llm
                progress_bar = st.progress(0, text="Generating report sections...")
                for i, title in enumerate(section_titles):
                    clean_title = re.sub(r'^\s*[\d\w]+\.\s*', '', title).strip()
                    st.write(f"-> Generating: {clean_title}")
                    response = section_chain.invoke({"section": clean_title})
                    full_report_list.append(f"## {title}\n\n{response.content}")
                    progress_bar.progress((i + 1) / len(section_titles), text=f"Generated: {clean_title}")
                final_report = f"# Your In-Depth Research Report: {report_topic}\n\n" + "\n\n".join(full_report_list)
                current_topic = report_topic
        
        if final_report:
            st.session_state.final_report = final_report
            st.session_state.report_topic = current_topic
            st.session_state.report_generated = True
            st.rerun()

    if st.session_state.report_generated:
        st.subheader("🎉 Your Detailed Research Report is Ready!")
        st.markdown(st.session_state.final_report)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            pdf_bytes = create_download_pdf(st.session_state.final_report, st.session_state.report_topic)
            if pdf_bytes:
                st.download_button(
                    label="⬇️ Download as PDF", 
                    data=pdf_bytes, 
                    file_name="report.pdf", 
                    mime="application/pdf",
                    use_container_width=True
                )
        
        with col2:
            docx_bytes = create_download_docx(st.session_state.final_report, st.session_state.report_topic)
            if docx_bytes:
                st.download_button(
                    label="⬇️ Download as DOCX",
                    data=docx_bytes,
                    file_name="report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        with col3:
            st.download_button(
                label="⬇️ Download as Markdown",
                data=st.session_state.final_report.encode('utf-8'),
                file_name="report.md",
                mime="text/markdown",
                use_container_width=True
            )

if __name__ == "__main__":
    main()